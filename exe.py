import docx
from openpyxl import load_workbook, Workbook
import openpyxl.utils.cell
import datetime
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
import os
from itertools import groupby

class referencias:
    def __init__(self, primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, url, not_italic, titulo_periodico, volume, acesso, local, data_conferencia, virgula, legislação, suporte, flag_datas):
        self.primeiro_nome = primeiro_nome
        self.segundo_nome = segundo_nome
        self.apelido = apelido
        self.referencia1 = referencia1
        self.referencia2 = referencia2
        self.titulo = titulo
        self.ano = ano
        self.url = url
        self.not_italic = not_italic
        self.titulo_periodico = titulo_periodico
        self.volume = volume
        self.acesso = acesso
        self.local = local
        self.data_conferencia = data_conferencia
        self.virgula = virgula
        self.legislação = legislação
        self.suporte = suporte
        self.flag_datas = flag_datas

lista = []
sorted_list = []
sorted_list2 = []

def run():

    data_file = 'refs.xlsx'
    wb = load_workbook(data_file)
    trabalho = input("Documento a validar: ")
    #trabalho = "ex.docx"
    doc = docx.Document(trabalho)
    referencias1 = []
    referencias2 = []
    titulos = []
    anos = []
    urls = []
    not_italic = []
    titulo_periodico = []
    volumes = []
    acessos = []
    locais = []
    datas_conferencias = []
    virgula = []
    legislacao = []
    suporte_arr = []
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    ids = 0

    def livros():

        livros = wb['livros']
        all_rows = list(livros.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "K" + str(i)
            if livros[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")

            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:

                primeiro_nome.append(f"{all_rows[n_row][1].value}")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                volume = ""
                if f"{all_rows[n_row][6].value}" != "None":
                    volume = f"{all_rows[n_row][6].value}"
                edição = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    edição = f"{all_rows[n_row][7].value}"
                    if edição == "1ª":
                        edição = ""
                local = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    local = f"{all_rows[n_row][8].value}"
                else:
                    local = "(s.l.)"
                editora = ""
                if f"{all_rows[n_row][9].value}" != "None":
                    editora = f"{all_rows[n_row][9].value}"
                else:
                    editora = "(s.n.)"
                while n > 1:
                    n_row += 1
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if volume != "":
                    referencia2 += " (" + volume + " volume)"
                if edição != "":
                    referencia2 += " (" + edição + " ed.)"
                referencia2 += "."
                if local != "":
                    referencia2 += " " + local
                if local != "" and editora != "":
                    referencia2 += ":"
                if editora != "":
                    referencia2 += " " + editora
                referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", "", ""))

            n_row += 1

    def livros_instituições():

        livros_instit = wb["livros publ. por instituições"]
        all_rows = list(livros_instit.rows)
        counter = 1

        for row in all_rows:
            nomes = []
            i = all_rows.index(row)+1
            column = "G" + str(i)
            if livros_instit[column].value == ".":
                nome = f"{all_rows[counter][0].value}"
                nomes.append(nome)
                ano = ""
                if f"{all_rows[counter][1].value}" != "None":
                    ano = f"{all_rows[counter][1].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[counter][2].value}"
                edição = ""
                if f"{all_rows[counter][3].value}" != "None":
                    edição = f"{all_rows[counter][3].value}"
                    if edição == "1ª":
                        edição = ""
                local = ""
                if f"{all_rows[counter][4].value}" != "None":
                    local = f"{all_rows[counter][4].value}"
                else:
                    local = "(s.l.)"
                editora = ""
                if f"{all_rows[counter][5].value}" != "None":
                    editora = f"{all_rows[counter][5].value}"
                else:
                    editora = "(s.n.)"

                referencia1 = ""
                referencia2 = ""

                referencia1 += nome + "."
                if edição != "":
                    referencia2 += " (" + edição + " ed.). "
                else:
                    referencia2 += "."
                if local != "":
                    referencia2 += " " + local
                if local != "" and editora != "":
                    referencia2 += ":"
                if editora != "":
                    referencia2 += " " + editora

                referencia2 += "."
                lista.append(referencias("", "", nomes, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", "", ""))

                counter += 1

    def ebooks():

        ebooks = wb['ebooks']
        all_rows = list(ebooks.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "I" + str(i)
            if ebooks[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                data = ""
                data_acesso = ""
                if f"{all_rows[n_row][6].value}" != "None":
                    data = f"{all_rows[n_row][6].value}"
                    dia = data[8] + data[9]
                    mês = data[5] + data[6]
                    mês = int(mês)
                    mês = meses[mês-1]
                    ano_acesso = data[0] + data[1] + data[2] + data[3]
                    data_acesso = dia + " de " + mês + " de " + ano_acesso
                url = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    url = f"{all_rows[n_row][7].value}"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i]
                        if primeiro_nome[i] != "":
                            referencia1 += ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                referencia2 += "."
                if data_acesso != "":
                    referencia2 += " Consultado a " + data_acesso + ". "
                if url != "":
                    referencia2 += " Disponível em "

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, url, "", "", "", "", "", "", "", "", "", ""))

            n_row += 1

    def artigos_fisicos():

        artigos_fisicos = wb['artigos físicos']
        all_rows = list(artigos_fisicos.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "N" + str(i)
            if artigos_fisicos[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                mês = ""
                if f"{all_rows[n_row][5].value}" != "None":
                    mês = meses[int(f"{all_rows[n_row][5].value}")-1]
                dia = ""
                if f"{all_rows[n_row][6].value}" != "None":
                    dia = f"{all_rows[n_row][6].value}"
                titulo = f"{all_rows[n_row][7].value}"
                titulo_publ = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    titulo_publ = f"{all_rows[n_row][8].value}"
                volume = ""
                if f"{all_rows[n_row][9].value}" != "None":
                    volume = f"{all_rows[n_row][9].value}"
                num = ""
                if f"{all_rows[n_row][10].value}" != "None":
                    num = f"{all_rows[n_row][10].value}"
                primeira_pag = ""
                if f"{all_rows[n_row][11].value}" != "None":
                    primeira_pag = f"{all_rows[n_row][11].value}"
                ultima_pag = ""
                if f"{all_rows[n_row][12].value}" != "None":
                    ultima_pag = f"{all_rows[n_row][12].value}"
                while n > 1:
                    n_row += 1
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if str(ano) != "" or mês != "" or dia != "":
                    referencia1 += " ("
                    if str(ano) != "":
                        referencia1 += str(ano)
                        if mês != "" or dia != "":
                            referencia1 += ", "
                    if mês != "":
                        referencia1 += mês
                    if mês != "" and dia != "":
                        referencia1 += " " + dia
                    else:
                        referencia1 += dia
                    referencia1 += ")"
                referencia1 += "."
                if num != "":
                    referencia2 += "(" + num + ")"
                if primeira_pag != "":
                    referencia2 += ", " + primeira_pag
                if primeira_pag != "" and ultima_pag != "":
                    referencia2 += "-"
                if ultima_pag != "":
                    referencia2 += ultima_pag
                referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, "", titulo, titulo_publ, volume, "", "", "", "", "", "", 1))

            n_row += 1

    def artigos_online():

        artigos_online = wb['artigos online']
        all_rows = list(artigos_online.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "N" + str(i)
            if artigos_online[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                titulo_publ = ""
                if f"{all_rows[n_row][6].value}" != "None":
                    titulo_publ = f"{all_rows[n_row][6].value}"
                volume = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    volume = f"{all_rows[n_row][7].value}"
                num = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    num = f"{all_rows[n_row][8].value}"
                primeira_pag = ""
                if f"{all_rows[n_row][9].value}" != "None":
                    primeira_pag = f"{all_rows[n_row][9].value}"
                ultima_pag = ""
                if f"{all_rows[n_row][10].value}" != "None":
                    ultima_pag = f"{all_rows[n_row][10].value}"
                data = ""
                data_acesso = ""
                if f"{all_rows[n_row][11].value}" != "None":
                    data = f"{all_rows[n_row][11].value}"
                    dia = data[8] + data[9]
                    mês = data[5] + data[6]
                    mês = int(mês)
                    mês = meses[mês-1]
                    ano_acesso = data[0] + data[1] + data[2] + data[3]
                    data_acesso = dia + " de " + mês + " de " + ano_acesso
                url = ""
                if f"{all_rows[n_row][12].value}" != "None":
                    url = f"{all_rows[n_row][12].value}"
                while n > 1:
                    n_row += 1
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if num != "":
                    if volume == "":
                        referencia2 += ", (" + num + ")"
                    else:
                        referencia2 += "(" + num + ")"
                if primeira_pag != "":
                    referencia2 += ", " + primeira_pag
                if primeira_pag != "" and ultima_pag != "":
                    referencia2 += "-"
                if ultima_pag != "":
                    referencia2 += ultima_pag
                referencia2 += ". "
                if data_acesso != "":
                    referencia2 += "Consultado a " + data_acesso + ". "
                if url != "":
                    referencia2 += "Disponível em "

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, url, titulo, titulo_publ, volume, "", "", "", "", "", "", ""))

            n_row += 1

    def teses():

        teses = wb['teses']
        all_rows = list(teses.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "L" + str(i)
            if teses[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                grau = f"{all_rows[n_row][6].value}"
                curso = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    curso = f"{all_rows[n_row][7].value}"
                universidade = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    universidade = f"{all_rows[n_row][8].value}"
                faculdade = ""
                if f"{all_rows[n_row][9].value}" != "None":
                    faculdade = f"{all_rows[n_row][9].value}"
                local = ""
                if f"{all_rows[n_row][10].value}" != "None":
                    local = f"{all_rows[n_row][10].value}"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                referencia2 += " ("
                if grau == "Licenciatura" or grau == "licenciatura":
                    referencia2 += "Monografia de Licenciatura"
                elif grau == "Mestrado" or grau == "mestrado":
                    referencia2 += "Tese de Mestrado"
                elif grau == "Doutoramento" or grau == "doutoramento":
                    referencia2 += "Tese de Doutoramento"
                referencia2 += " não editada"
                if curso != "":
                    referencia2 += ", " + curso
                referencia2 += "). "
                if universidade != "":
                    referencia2 += universidade
                    if faculdade != "" or local != "":
                        referencia2 += ", "
                if faculdade != "":
                    referencia2 += faculdade
                if local != "":
                    if faculdade != "":
                        referencia2 += ", " + local 
                    else:
                        referencia2 += local
                else:
                    referencia2 += local
                referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", "", ""))

            n_row += 1

    def teses_bd():

        teses_bd = wb['teses via base de dados']
        all_rows = list(teses_bd.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "K" + str(i)
            if teses_bd[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                grau = f"{all_rows[n_row][6].value}"
                bd = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    bd = f"{all_rows[n_row][7].value}"
                url = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    url = f"{all_rows[n_row][8].value}"
                acesso = ""
                if f"{all_rows[n_row][9].value}" != "None":
                    acesso = f"{all_rows[n_row][9].value}"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                referencia2 += " ("
                if grau == "Mestrado" or grau == "mestrado":
                    referencia2 += "Tese de Mestrado). "
                elif grau == "Doutoramento" or grau == "doutoramento":
                    referencia2 += "Tese de Doutoramento). "
                if bd != "" or url != "":
                    referencia2 += "Disponível em "
                if bd != "":
                    referencia2 += bd
                    if url != "":
                        referencia2 += " em "
                    if url == "" and acesso == "":
                        referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, url, "", "", "", acesso, "", "", "", "", "", ""))

            n_row += 1

    def relat_tecnico():

        relat_tecnicos = wb['relat. técnico ou de pesquisa']
        all_rows = list(relat_tecnicos.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "J" + str(i)
            if relat_tecnicos[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                n_relatorio = ""
                if f"{all_rows[n_row][6].value}" != "None":
                    n_relatorio = f"{all_rows[n_row][6].value}"
                local = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    local = f"{all_rows[n_row][7].value}"
                uni = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    uni = f"{all_rows[n_row][8].value}"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if n_relatorio != "":
                    referencia2 += " (Relatório nº " + n_relatorio + ")."
                if local != "":
                    referencia2 += " " + local
                    if uni != "":
                        referencia2 += ":"
                if uni != "":
                    referencia2 += " " + uni
                referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", "", ""))

            n_row += 1

    def relat_tecnico_online():

        relat_tecnicos_online = wb['o mesmo mas online']
        all_rows = list(relat_tecnicos_online.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "J" + str(i)
            if relat_tecnicos_online[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                n_relatorio = ""
                if f"{all_rows[n_row][6].value}" != "None":
                    n_relatorio = f"{all_rows[n_row][6].value}"
                uni = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    uni = f"{all_rows[n_row][7].value}"
                url = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    url = f"{all_rows[n_row][8].value}"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if n_relatorio != "":
                    referencia2 += " (Relatório nº " + n_relatorio + ")."
                if uni != "":
                    referencia2 += " Disponível via " + uni
                    if url != "":
                        referencia2 += " em "
                    else:
                        referencia2 += "."
                if url != "":
                    if url != "" and uni == "":
                        referencia2 += " Disponível em "

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, url, "", "", "", "", "", "", "", "", "", ""))

            n_row += 1

    def comunicação_congresso():

        com_congresso = wb['comun. conferência e congresso']
        all_rows = list(com_congresso.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "J" + str(i)
            if com_congresso[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []

            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                mês = ""
                if f"{all_rows[n_row][5].value}" != "None":
                    mês = f"{all_rows[n_row][5].value}"
                    mês = int(mês)
                    mês = meses[mês-1]
                titulo = f"{all_rows[n_row][6].value}"
                congresso = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    congresso = f"{all_rows[n_row][7].value}"
                local = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    local = f"{all_rows[n_row][8].value}"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if ano != "" or mês != "":
                    referencia1 += " (" 
                    if ano != "":
                        referencia1 += str(ano)
                        if mês != "":
                            referencia1 += ", "
                    if mês != "":
                        referencia1 += mês
                    referencia1 += ")."
                if congresso != "":
                    referencia2 += ". Apresentada em " + congresso
                    if local != "":
                        referencia2 += ","
                if local != "":
                    referencia2 += " " + local
                referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", "", 1))

            n_row += 1

    def atas():

        atas = wb['atas conferência e congresso']
        all_rows = list(atas.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "K" + str(i)
            if atas[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []

            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[n_row][5].value}"
                primeira_pag = ""
                if f"{all_rows[n_row][6].value}" != "None":
                    primeira_pag = f"{all_rows[n_row][6].value}"
                segunda_pag = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    segunda_pag = f"{all_rows[n_row][7].value}"
                local = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    local = f"{all_rows[n_row][8].value}"
                editora = ""
                if f"{all_rows[n_row][9].value}" != "None":
                    editora = f"{all_rows[n_row][9].value}"
                else:
                    editora = "[s.n.]"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if primeira_pag != "":
                    referencia2 += " ("
                    if primeira_pag == segunda_pag:
                        referencia2 += primeira_pag
                    else:
                        referencia2 += primeira_pag + "-" + segunda_pag
                    referencia2 += ")."
                else:
                    referencia2 += "."
                if local != "":
                    referencia2 += " " + local
                    if editora != "":
                        referencia2 += ":"
                if editora != "":
                    referencia2 += " " + editora
                if local != "" or editora != "":
                    referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", "", ""))

            n_row += 1

    def artigos_atas():

        arts_atas = wb['art. num livro de atas de congr']
        all_rows = list(arts_atas.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "R" + str(i)
            if arts_atas[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []

            if n >= 1:
                if f"{all_rows[n_row][1].value}" != "None":
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                else:
                    primeiro_nome.append("")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                ano = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    ano = f"{all_rows[n_row][4].value}"
                    ano = int(ano)
                mês = ""
                if f"{all_rows[n_row][5].value}" != "None":
                    mês = f"{all_rows[n_row][5].value}"
                    mês = meses[mês-1]
                titulo = f"{all_rows[n_row][6].value}"
                ata = ""
                if f"{all_rows[n_row][7].value}" != "":
                    ata = f"{all_rows[n_row][7].value}"
                local_congresso = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    local_congresso = f"{all_rows[n_row][8].value}"
                ano_congresso = ""
                if f"{all_rows[n_row][9].value}" != "None":
                    ano_congresso = f"{all_rows[n_row][9].value}"
                mês_congresso = ""
                if f"{all_rows[n_row][10].value}" != "None":
                    mês_congresso = f"{all_rows[n_row][10].value}"
                    mês_congresso = int(mês_congresso)
                    mês_congresso = meses[mês_congresso-1]
                primeiro_dia_congresso = ""
                if f"{all_rows[n_row][11].value}" != "None":
                    primeiro_dia_congresso = f"{all_rows[n_row][11].value}"
                ultimo_dia_congresso = ""
                if f"{all_rows[n_row][12].value}" != "None":
                    ultimo_dia_congresso = f"{all_rows[n_row][12].value}"
                primeira_pag = ""
                if f"{all_rows[n_row][13].value}" != "None":
                    primeira_pag = f"{all_rows[n_row][13].value}"
                ultima_pag = ""
                if f"{all_rows[n_row][14].value}" != "None":
                    ultima_pag = f"{all_rows[n_row][14].value}"
                local = ""
                if f"{all_rows[n_row][15].value}" != "None":
                    local = f"{all_rows[n_row][15].value}"
                organização = ""
                if f"{all_rows[n_row][16].value}" != "None":
                    organização = f"{all_rows[n_row][16].value}"
                while n > 1:
                    n_row += 1
                    if f"{all_rows[n_row][1].value}" != "None":
                        primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    else:
                        primeiro_nome.append("")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if ano != "" or mês != "":
                    referencia1 += " ("
                    if ano != "":
                        referencia1 += str(ano)
                        if mês != "":
                            referencia1 += ", "
                    if mês != "":
                        referencia1 += mês
                    referencia1 += ")"
                referencia1 += "."
                if titulo != "":
                    referencia1 += " " + titulo + "."
                if ata != "":
                    referencia1 += " In"
                datas_congresso = ""
                if primeiro_dia_congresso != "":
                    if primeiro_dia_congresso == ultimo_dia_congresso:
                        datas_congresso += primeiro_dia_congresso
                    else:
                        datas_congresso += primeiro_dia_congresso + "-" + ultimo_dia_congresso
                    datas_congresso += " "
                    if mês_congresso != "" or ano_congresso != "":
                        datas_congresso += "de "
                if mês_congresso != "":
                    datas_congresso += mês_congresso
                    if ano_congresso != "":
                        datas_congresso += " de "
                if ano_congresso != "":
                    datas_congresso += ano_congresso
                if primeira_pag != "" or local != "" or organização != "":
                    referencia2 += " "
                if primeira_pag != "":
                    referencia2 += "("
                    if primeira_pag == ultima_pag:
                        referencia2 += "p. " + primeira_pag
                    else:
                        referencia2 += "pp. " + primeira_pag + "-" + ultima_pag
                    referencia2 += ")"
                referencia2 += ". "
                if local != "":
                    referencia2 += local
                    if organização != "":
                        referencia2 += ": "
                    else:
                        referencia2 += "."
                if organização != "":
                    referencia2 += organização + "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, ata, ano, "", "", "", "", "", local_congresso, datas_congresso, ata, "", "", 1))

            n_row += 1

    def legislação():

        legislação = wb["legislação portuguesa"]
        all_rows = list(legislação.rows)
        counter = 1

        for row in all_rows:
            titulo = []
            i = all_rows.index(row)+1
            column = "I" + str(i)
            if legislação[column].value == ".":
                t = f"{all_rows[counter][0].value}"
                titulo.append(t)
                n_lei = f"{all_rows[counter][1].value}"
                instituição = ""
                if f"{all_rows[counter][2].value}" != "None":
                    instituição = f"{all_rows[counter][2].value}"
                ano = ""
                if f"{all_rows[counter][3].value}" != "None":
                    ano = f"{all_rows[counter][3].value}"
                    ano = int(ano)
                publicação = ""
                if f"{all_rows[counter][4].value}" != "None":
                    publicação = f"{all_rows[counter][4].value}"
                serie = ""
                if f"{all_rows[counter][5].value}" != "None":
                    serie = f"{all_rows[counter][5].value}"
                n_publicação = ""
                if f"{all_rows[counter][6].value}" != "":
                    n_publicação = f"{all_rows[counter][6].value}"
                url = ""
                if f"{all_rows[counter][7].value}" != "None":
                    url = f"{all_rows[counter][7].value}"

                referencia1 = ""
                referencia2 = ""
                italico = ""

                referencia1 += t + " " + n_lei
                if instituição != "":
                    referencia1 += " do " + instituição
                referencia1 += "."                
                if publicação != "":
                    italico += publicação
                    if serie != "" or n_publicação != "":
                        italico += ":"
                    else:
                        italico += "."
                    italico += " "
                if serie != "":
                    italico += serie
                    if n_publicação != "":
                        italico += ","
                    else:
                        italico += "."
                    italico += " "
                if n_publicação != "":
                    italico += n_publicação + "."
                referencia2 += " "

                lista.append(referencias("", "", titulo, referencia1, referencia2, titulo, ano, url, "", italico, "", "", "", "", titulo, titulo, "", 1))

                counter += 1

    def audiovisuais():

        av = wb["audiovisuais"]
        all_rows = list(av.rows)
        counter = 1

        for row in all_rows:
            pn = []
            sn = []
            a = []
            i = all_rows.index(row)+1
            column = "J" + str(i)
            if av[column].value == ".":
                primeiro_nome = "" 
                if f"{all_rows[counter][0].value}" != "None":
                    primeiro_nome = f"{all_rows[counter][0].value}"
                    if f"{all_rows[counter][2].value}" != "None":
                        primeiro_nome = primeiro_nome[0] + "."
                pn.append(primeiro_nome)
                segundo_nome = ""
                if f"{all_rows[counter][1].value}" != "None":
                    segundo_nome = f"{all_rows[counter][1].value}"
                    segundo_nome = segundo_nome[0] + "."
                sn.append(segundo_nome)
                apelido = ""
                if f"{all_rows[counter][2].value}" != "None":
                    apelido = f"{all_rows[counter][2].value}"
                a.append(apelido)
                função = ""
                if f"{all_rows[counter][3].value}" != "None":
                    função = f"{all_rows[counter][3].value}"
                ano = ""
                if f"{all_rows[counter][4].value}" != "None":
                    ano = f"{all_rows[counter][4].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[counter][5].value}"
                suporte = ""
                if f"{all_rows[counter][6].value}" != "None":
                    suporte = f"{all_rows[counter][6].value}"
                local = ""
                if f"{all_rows[counter][7].value}" != "None":
                    local = f"{all_rows[counter][7].value}"
                else:
                    local = "[s.l.]"
                distribuidor = ""
                if f"{all_rows[counter][8].value}" != "None":
                    distribuidor = f"{all_rows[counter][8].value}"
                else:
                    distribuidor = "[s.n.]"

                referencia1 = ""
                referencia2 = ""

                referencia1 += apelido
                if apelido != "" and primeiro_nome != "":
                    referencia1 += ", "
                if primeiro_nome != "":
                    referencia1 += primeiro_nome
                if segundo_nome != "":
                    referencia1 += " " + segundo_nome
                if função != "":
                    referencia1 += " (" + função + ")."
                if suporte != "":
                    referencia2 += " [" + suporte + "]"
                referencia2 += ". "
                if local != "":
                    referencia2 += local
                    if distribuidor != "":
                        referencia2 += ": "
                    else:
                        referencia2 += "."
                if distribuidor != "":
                    referencia2 += distribuidor + "."

                if apelido != "":
                    lista.append(referencias(pn, sn, a, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", titulo, ""))
                else:
                    lista.append(referencias("", "", pn, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", titulo, ""))

                counter += 1

    def web():

        endereços_web = wb["endereços web"]
        all_rows = list(endereços_web.rows)
        counter = 1

        for row in all_rows:
            n = []
            i = all_rows.index(row)+1
            column = "F" + str(i)
            if endereços_web[column].value == ".":
                nome = f"{all_rows[counter][0].value}"
                n.append(nome)
                ano = ""
                if f"{all_rows[counter][1].value}" != "None":
                    ano = f"{all_rows[counter][1].value}"
                    ano = int(ano)
                else:
                    ano = 0
                titulo = f"{all_rows[counter][2].value}"
                data = f"{all_rows[counter][3].value}"
                a = data[:4]
                m = int(data[5:7])
                m = meses[m-1]
                d = data[8:10]
                data_acesso = d + " de " + m + " de " + a
                url = ""
                if f"{all_rows[counter][4].value}"  != "None":
                    url = f"{all_rows[counter][4].value}"

                referencia1 = ""
                referencia2 = ""

                referencia1 += nome + "."
                referencia2 += "."
                if data_acesso != "":
                    referencia2 += " Consultado a " + data_acesso + "."
                if url != "":
                    referencia2 += " Disponível em "

                lista.append(referencias("", "", n, referencia1, referencia2, titulo, ano, url, "", "", "", "", "", "", "", "", "", ""))

                counter += 1

    def patentes():

        patents = wb['patentes']
        all_rows = list(patents.rows)
        counter = 1

        for row in all_rows:
            i = all_rows.index(row)+1
            column = "J" + str(i)
            if patents[column].value == ".":
                counter += 1
        
        n_row = 1

        while n_row < counter:
            n = 0
            if f"{all_rows[n_row][0].value}" != "None":
                n = int(f"{all_rows[n_row][0].value}")
            primeiro_nome = []
            segundo_nome = []
            apelido = []
            if n >= 1:
                primeiro_nome.append(f"{all_rows[n_row][1].value}")
                if f"{all_rows[n_row][2].value}" != "None":
                    segundo_nome.append(f"{all_rows[n_row][2].value}")
                else:
                    segundo_nome.append("")
                apelido.append(f"{all_rows[n_row][3].value}")
                função = ""
                if f"{all_rows[n_row][4].value}" != "None":
                    função = f"{all_rows[n_row][4].value}"
                ano = ""
                if f"{all_rows[n_row][5].value}" != "None":
                    ano = f"{all_rows[n_row][5].value}"
                    ano = int(ano)
                else:
                    anos.append(0)
                titulo = f"{all_rows[n_row][6].value}"
                local = ""
                if f"{all_rows[n_row][7].value}" != "None":
                    local = f"{all_rows[n_row][7].value}"
                else:
                    local = "[s.l.]"
                nome_editor = ""
                if f"{all_rows[n_row][8].value}" != "None":
                    nome_editor = f"{all_rows[n_row][8].value}"
                else:
                    nome_editor = "[s.n.]"
                while n > 1:
                    n_row += 1
                    primeiro_nome.append(f"{all_rows[n_row][1].value}")
                    if f"{all_rows[n_row][2].value}" != "None":
                        segundo_nome.append(f"{all_rows[n_row][2].value}")
                    else:
                        segundo_nome.append("")
                    apelido.append(f"{all_rows[n_row][3].value}")
                    n -= 1

                referencia1 = ""
                referencia2 = ""

                autores = len(apelido)
                if autores < 3:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if referencia1 != "":
                            referencia1 += ", & "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1
                else:
                    for nome in apelido:
                        i = apelido.index(nome)
                        if i == autores-1:
                            if referencia1 != "":
                                referencia1 += ", & "
                        else:
                            if referencia1 != "":
                                referencia1 += ", "
                        referencia1 += apelido[i] + ", " + primeiro_nome[i][0] + "."
                        if segundo_nome[i] != "":
                            referencia1 += " " + segundo_nome[i][0] + "."
                        i += 1

                if função != "":
                    referencia1 += " (" + função + ")"
                    if ano != "":
                        referencia1 += "."
                if local != "" or nome_editor != "":
                    referencia2 += ","
                if local != "":
                    referencia2 += " " + local
                    if nome_editor != "":
                        referencia2 += ":"
                if nome_editor != "":
                    referencia2 += " " + nome_editor
                referencia2 += "."

                lista.append(referencias(primeiro_nome, segundo_nome, apelido, referencia1, referencia2, titulo, ano, "", "", "", "", "", "", "", "", "", "", ""))

            n_row += 1

    def ordem_alfabetica():

        global sorted_list
        global lista

        sorted_list = sorted(lista, key=lambda x: x.apelido)
        
        return sorted_list
    
    def ordem_anos():

        global sorted_list
        global sorted_list2

        # Group objects by their "apelido" attribute
        grouped_objects = [list(group) for _, group in groupby(sorted_list, key=lambda x: x.apelido)]
        
        for group in grouped_objects:
        # Create a dictionary to store objects with the same "ano"
            objects_by_year = {}
            
            for item in group:
                if item.ano not in objects_by_year:
                    objects_by_year[item.ano] = [item]
                else:
                    objects_by_year[item.ano].append(item)
            
            # Sort objects within each "ano" group by "titulo"
            for year_group in objects_by_year.values():
                year_group.sort(key=lambda x: x.titulo)
            
            # Add a letter suffix to differentiate objects with the same "ano"
            for year_group in objects_by_year.values():
                if len(year_group) > 1:
                    letter_counter = {}
                    for item in year_group:
                        if item.ano in letter_counter:
                            letter_counter[item.ano] += 1
                        else:
                            letter_counter[item.ano] = 1
                        
                        letter = chr(ord('a') + letter_counter[item.ano] - 1)
                        item.ano = f"{item.ano}{letter}"
            
            # Flatten and update the group
            group[:] = [item for year_group in objects_by_year.values() for item in year_group]
        
        # Flatten the grouped and sorted list back
        sorted_list2 = [item for group in grouped_objects for item in group]
        return sorted_list2

    

    livros() 
    livros_instituições()
    ebooks()
    artigos_fisicos()
    artigos_online()
    teses()
    teses_bd()
    relat_tecnico()
    relat_tecnico_online()
    comunicação_congresso()
    atas()
    artigos_atas()
    legislação()
    audiovisuais()
    web()
    patentes()
    ordem_alfabetica()
    ordem_anos()

    for ref in sorted_list2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(ref.referencia1)
        if ref.flag_datas == "":
            if ref.ano == 0:
                p.add_run(" (s.d.).")
            else:
                p.add_run(" (" + str(ref.ano) + ").")
        flag = 0
        flag2 = 0
        flag3 = 0
        flag4 = 0
        p.add_run(" ")
        for i in sorted_list2:
            if ref.titulo == i.not_italic:
                flag = 1
        if flag == 0 and ref.titulo != ref.legislação:
            p.add_run(ref.titulo).italic = True
        else:
            for i in sorted_list2:
                if ref.titulo == i.virgula:
                    flag2 = 1
            if flag2 == 1:
                if ref.local != "" or ref.data_conferencia != "":
                    p.add_run(" " + ref.titulo + ",")
            else:
                for i in sorted_list2:
                    if ref.titulo == i.legislação:
                        flag3 = 1
                if flag3 == 0 and ref.titulo != ref.legislação:
                    for i in sorted_list2:
                        if ref.titulo == i.suporte:
                            flag4 = 1
                    if flag4 == 0:
                        p.add_run(ref.titulo)
                        p.add_run(".")
                    else:
                        p.add_run(ref.titulo)
        if ref.local != "":
            p.add_run(", ")
            p.add_run(ref.local).italic = True
            if ref.data_conferencia != "":
                p.add_run(",")
        if ref.data_conferencia != "":
            p.add_run(" ")
            p.add_run(ref.data_conferencia).italic = True
        if ref.titulo_periodico != "":
            if ref.legislação == "":
                p.add_run(" ")
                p.add_run(ref.titulo_periodico).italic = True
            else:
                p.add_run(ref.titulo_periodico).italic = True
        if ref.volume != "":
            if ref.titulo_periodico != "":
                p.add_run(",")
            p.add_run(" ")
            p.add_run(ref.volume).italic = True
        p.add_run(ref.referencia2)
        if ref.url != "":
            run = p.add_run(ref.url)
            run.font.underline = True
            font = run.font
            font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        if ref.acesso != "":
            p.add_run(" (" + ref.acesso + ").")
        
            

    doc.save(trabalho)
    print("")
    print("All done!")
    print("")
    os.system(trabalho)

run()